from __future__ import annotations

import asyncio
import io
import os
import re
import shutil
from dataclasses import dataclass

from app.core.errors import AppError

SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "image/png": "png", "image/jpeg": "jpg", "image/tiff": "tiff",
}


@dataclass
class ExtractedDocument:
    text: str
    page_count: int | None = None


@dataclass
class Chunk:
    index: int
    content: str
    page_number: int | None = None


class DocumentService:
    def validate(self, filename: str, content_type: str, content: bytes, max_bytes: int) -> str:
        if not content or len(content) > max_bytes:
            raise AppError("invalid_document_size", f"Document must be between 1 and {max_bytes} bytes", status_code=413)
        document_type = SUPPORTED_TYPES.get(content_type.split(";", 1)[0].lower())
        if not document_type:
            raise AppError("unsupported_document_type", "Only PDF, DOCX, TXT, PNG, JPEG, and TIFF are supported", status_code=415)
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        allowed_extensions = {"jpg", "jpeg"} if document_type == "jpg" else {document_type}
        if extension not in allowed_extensions:
            raise AppError("document_type_mismatch", "Filename extension does not match the content type", status_code=415)
        signatures = {
            "pdf": content.startswith(b"%PDF-"), "docx": content.startswith(b"PK"),
            "png": content.startswith(b"\x89PNG\r\n\x1a\n"),
            "jpg": content.startswith(b"\xff\xd8\xff"),
            "tiff": content.startswith((b"II*\x00", b"MM\x00*")),
        }
        if document_type == "txt":
            try:
                content.decode("utf-8-sig")
            except UnicodeDecodeError as exc:
                raise AppError("invalid_text_encoding", "Text documents must use UTF-8 encoding", status_code=415) from exc
        elif not signatures.get(document_type, False):
            raise AppError("invalid_document_signature", "File content does not match its declared type", status_code=415)
        return document_type

    async def extract(self, content: bytes, document_type: str) -> ExtractedDocument:
        return await asyncio.to_thread(self._extract_sync, content, document_type)

    def _extract_sync(self, content: bytes, document_type: str) -> ExtractedDocument:
        if document_type == "pdf":
            return self._extract_pdf(content)
        if document_type == "docx":
            from docx import Document
            doc = Document(io.BytesIO(content))
            return ExtractedDocument("\n".join(p.text for p in doc.paragraphs if p.text.strip()))
        if document_type == "txt":
            return ExtractedDocument(content.decode("utf-8-sig"))
        from PIL import Image
        return ExtractedDocument(self._ocr_image(Image.open(io.BytesIO(content))))

    def _extract_pdf(self, content: bytes) -> ExtractedDocument:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(content))
        page_text = [(page.extract_text() or "").strip() for page in reader.pages]
        empty_pages = [index for index, text in enumerate(page_text) if not text]
        if empty_pages:
            try:
                import fitz
            except ImportError as exc:
                raise RuntimeError("Scanned PDF support is unavailable because PyMuPDF is not installed") from exc

            pdf = fitz.open(stream=content, filetype="pdf")
            from PIL import Image
            try:
                for index in empty_pages:
                    pixmap = pdf[index].get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                    page_text[index] = self._ocr_image(image).strip()
            finally:
                pdf.close()
        return ExtractedDocument("\n\n".join(page_text), len(reader.pages))

    @staticmethod
    def _ocr_image(image) -> str:
        import pytesseract
        if not shutil.which("tesseract") and os.name == "nt":
            common_install = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            if os.path.isfile(common_install):
                pytesseract.pytesseract.tesseract_cmd = common_install
        try:
            return pytesseract.image_to_string(image)
        except pytesseract.TesseractNotFoundError as exc:
            raise RuntimeError(
                "This document contains scanned images. Install Tesseract OCR on Windows, "
                "then restart the EVA worker."
            ) from exc

    def clean(self, text: str) -> str:
        text = text.replace("\x00", "")
        return re.sub(r"[ \t]+", " ", re.sub(r"\n{3,}", "\n\n", text)).strip()

    def chunk(self, text: str, target_chars: int = 1600, overlap_chars: int = 200) -> list[Chunk]:
        if not text:
            return []
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        chunks, current = [], ""
        for paragraph in paragraphs:
            pieces = [paragraph[i:i + target_chars] for i in range(0, len(paragraph), target_chars)]
            for piece in pieces:
                if current and len(current) + len(piece) + 2 > target_chars:
                    chunks.append(Chunk(len(chunks), current))
                    current = current[-overlap_chars:] + "\n\n" + piece
                else:
                    current = f"{current}\n\n{piece}".strip()
        if current:
            chunks.append(Chunk(len(chunks), current))
        return chunks
